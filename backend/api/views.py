from io import BytesIO
from django.db.models import Sum
from django.urls import reverse
from django_filters.rest_framework import DjangoFilterBackend
from django.http import FileResponse, HttpResponseRedirect
from django.shortcuts import get_object_or_404
from rest_framework import response
from docx import Document
from rest_framework import status, viewsets
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated
from django.db.models.functions import Lower
from django.contrib.auth import get_user_model
from djoser.views import UserViewSet as DjoserUserViewSet
from rest_framework.response import Response
from rest_framework import permissions
from django.shortcuts import redirect
from hashids import Hashids
from django.conf import settings

from . import serializers as myserializers
from food.models import (
    Recipe,
    Tag,
    Ingredient,
    Favorite,
    ShoppingCart,
    RecipeIngredient,
)
from users.models import Sub
from api import (
    paginators as tools_paginators,
    filters as tools_filters,
    permissions as tools_permissions,
)


hashids = Hashids(salt=settings.SECRET_KEY, min_length=5)

User = get_user_model()


class UserViewSet(DjoserUserViewSet):
    """
    Кастомное представление для пользователей.
    """

    pagination_class = tools_paginators.Paginator

    def get_serializer_class(self):
        if self.action == "list":
            return myserializers.UserSerializer
        elif self.action == "retrieve":
            return myserializers.UserSerializer
        return super().get_serializer_class()

    def get_permissions(self):
        if self.action == "list":
            return [permissions.AllowAny()]
        elif self.action == "retrieve":
            return [permissions.AllowAny()]
        return super().get_permissions()

    @action(
        methods=["get"],
        detail=False,
        url_path="me",
        permission_classes=[IsAuthenticated],
    )
    def me(self, request, *args, **kwargs):
        return super().me(request, *args, **kwargs)

    @action(
        detail=False,
        methods=["put"],
        permission_classes=[IsAuthenticated],
        url_path="me/avatar",
    )
    def avatar(self, request, *args, **kwargs):
        serializer = myserializers.AvatarSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        request.user.avatar = serializer.validated_data["avatar"]
        request.user.save()

        response_serializer = myserializers.AvatarSerializer(
            request.user, context={"request": request}
        )
        return Response(response_serializer.data, status=status.HTTP_200_OK)

    @avatar.mapping.delete
    def avatar_delete(self, request, *args, **kwargs):
        if request.user.avatar:
            request.user.avatar.delete()
            request.user.avatar = None
            request.user.save()
        return Response(status=status.HTTP_204_NO_CONTENT)

    @action(
        detail=True,
        methods=["post"],
        permission_classes=[permissions.IsAuthenticated],
        url_path="subscribe",
    )
    def subscribe(self, request, *args, **kwargs):
        author = self.get_object()
        user = request.user

        serializer = myserializers.SubscriptionCreateSerializer(
            data={"user": user.id, "author": author.id},
            context={"request": request},
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()

        return response.Response(
            serializer.data,
            status=status.HTTP_201_CREATED
        )

    @subscribe.mapping.delete
    def unsubscribe(self, request, *args, **kwargs):
        author = self.get_object()
        user = request.user

        deleted_count, _ = Sub.objects.filter(
            user=user,
            author=author
        ).delete()

        if not deleted_count:
            return response.Response(
                {"errors": "Подписка не найдена."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        return response.Response(status=status.HTTP_204_NO_CONTENT)

    @action(
        detail=False,
        methods=["get"],
        permission_classes=[permissions.IsAuthenticated],
        url_path="subscriptions",
    )
    def subscriptions(self, request, *args, **kwargs):
        user = request.user
        queryset = User.objects.filter(users_subscribers__user=user)
        pages = self.paginate_queryset(queryset)
        serializer = myserializers.SubscribeSerializer(
            pages, many=True, context={"request": request}
        )
        return self.get_paginated_response(serializer.data)


class RecipeViewSet(viewsets.ModelViewSet):
    queryset = Recipe.objects.all()
    permission_classes = [
        permissions.IsAuthenticatedOrReadOnly,
        tools_permissions.IsAuthorOrReadOnlyPermission,
    ]
    pagination_class = tools_paginators.Paginator
    filter_backends = [DjangoFilterBackend]
    filterset_class = tools_filters.RecipeFilter

    def get_serializer_class(self):
        if self.action in ["list", "retrieve"]:
            return myserializers.RecipeReadSerializer
        return myserializers.RecipeWriteSerializer

    def perform_create(self, serializer):
        recipe = serializer.save()

        recipe.short_code = hashids.encode(recipe.id)
        recipe.save()

    @action(detail=False, methods=["get"], url_path="r/(?P<short_code>[^/.]+)")
    def redirect_by_short_code(self, request, short_code=None):
        recipe = get_object_or_404(Recipe, short_code=short_code)
        return redirect(f"/recipes/{recipe.id}/")

    @action(
        detail=True,
        methods=['get'],
        url_path='get-link'
    )
    def link(self, request, *args, **kwargs):
        recipe = self.get_object()
        base_url = request.build_absolute_uri('/')[:-1]
        short_url = f'{base_url}/api/recipes/r/{recipe.short_code}'
        return response.Response(
            {'short-link': short_url}, status=status.HTTP_200_OK
        )

    @action(
        detail=True,
        methods=["post"],
        url_name="shopping_cart",
        permission_classes=[permissions.IsAuthenticated],
    )
    def shopping_cart(self, request, pk=None):
        recipe = get_object_or_404(Recipe, id=pk)

        serializer = myserializers.ShoppingCartSerializer(
            data={"user": request.user.id, "recipe": recipe.id},
            context={"request": request},
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()

        return response.Response(
            serializer.data,
            status=status.HTTP_201_CREATED
        )

    @shopping_cart.mapping.delete
    def remove_from_shopping_cart(self, request, pk=None):
        recipe = get_object_or_404(Recipe, pk=pk)
        deleted_count, _ = ShoppingCart.objects.filter(
            user=request.user, recipe=recipe
        ).delete()

        if not deleted_count:
            return response.Response(
                {"errors": "Рецепт не найден в корзине"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        return response.Response(status=status.HTTP_204_NO_CONTENT)

    @action(
        detail=True,
        methods=["post"],
        url_name="favorite",
        permission_classes=[permissions.IsAuthenticated],
    )
    def favorite(self, request, pk=None):
        recipe = get_object_or_404(Recipe, id=pk)

        serializer = myserializers.FavoriteSerializer(
            data={"user": request.user.id, "recipe": recipe.id},
            context={"request": request},
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()

        recipe_serializer = myserializers.RecipeShortSerializer(
            recipe,
            context={"request": request}
        )
        return response.Response(
            recipe_serializer.data,
            status=status.HTTP_201_CREATED
        )

    @favorite.mapping.delete
    def remove_from_favorite(self, request, pk=None):
        recipe = get_object_or_404(Recipe, pk=pk)
        deleted_count, _ = Favorite.objects.filter(
            user=request.user, recipe=recipe
        ).delete()

        if not deleted_count:
            return response.Response(
                {"errors": "Рецепт не найден в избранном"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        return response.Response(status=status.HTTP_204_NO_CONTENT)

    @action(
        detail=False,
        methods=["get"],
        url_path="download_shopping_cart",
        permission_classes=[permissions.IsAuthenticated],
    )
    def download_shopping_cart(self, request):
        document = Document()
        document.add_heading("Список покупок", level=1)

        recipes_shop = request.user.shopping_cart.all()
        ingredients = (
            RecipeIngredient.objects.filter(
                recipe__in=recipes_shop.values_list("recipe", flat=True)
            )
            .values("ingredient__name", "ingredient__measurement_unit")
            .annotate(total_amount=Sum("amount"))
            .order_by("ingredient__name")
        )

        if ingredients:
            table = document.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "№"
            hdr_cells[1].text = "Ингредиент"
            hdr_cells[2].text = "Количество"

            for idx, ing in enumerate(ingredients, start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = ing["ingredient__name"]
                row_cells[2].text = (
                    f"{ing['total_amount']} "
                    f"{ing['ingredient__measurement_unit']}"
                )
        else:
            document.add_paragraph("Список покупок пуст!")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        content_type = (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
        return FileResponse(
            buffer,
            as_attachment=True,
            filename="shopping_list.docx",
            content_type=content_type,
        )


class TagViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Получение списка тэгов.
    Создание и редактирование только в админке.
    """

    queryset = Tag.objects.all()
    serializer_class = myserializers.TagSerializer
    pagination_class = None


class IngredientViewSet(viewsets.ReadOnlyModelViewSet):
    """
    Получение списка ингредиентов.
    Создание и редактирование только в админке.
    """

    queryset = Ingredient.objects.all().order_by(Lower("name"))
    serializer_class = myserializers.IngredientSerializer
    filter_backends = (DjangoFilterBackend,)
    filterset_class = tools_filters.IngredientSearchFilter
    pagination_class = None


def redirect_to_recipe(request, short_code):
    recipe = get_object_or_404(Recipe, short_code=short_code)
    url = request.build_absolute_uri(
        reverse("recipes-detail", kwargs={"pk": recipe.pk})
    )
    return HttpResponseRedirect(url)
